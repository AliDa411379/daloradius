from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2F5496')
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if r % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')
    return table

# ============ TITLE ============
title = doc.add_heading('DaloRADIUS - ERP Integration API', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Task Assignment Sheet')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(47, 84, 150)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Date: 2026-03-09  |  Team Size: 3 Developers').font.size = Pt(10)

doc.add_paragraph()

# ============ TEAM ROLES ============
doc.add_heading('Team Roles', level=1)

add_styled_table(doc,
    ['Role', 'Developer', 'Scope', 'Technology'],
    [
        ['ERP Frontend', 'Wael',
         'ERP Dashboard UI - consumes ALL APIs built by Hisham & Laith.\n'
         'Builds login page, pages, tables, charts, action buttons.\n'
         'Calls POST /erp_login first to get token,\n'
         'then sends token in Authorization header on all requests.',
         'ERP Framework (frontend)\nHTTP/REST client\nCharts library'],
        ['DaloRADIUS API (GET)', 'Hisham',
         'Builds all READ/GET endpoints.\n'
         'User status, user details, usage data, plan profiles, dashboard summary.\n'
         'All endpoints require valid token.',
         'PHP / MySQL\nDaloRADIUS codebase'],
        ['DaloRADIUS API (DB + POST + Auth)', 'Laith',
         'Builds login/token auth system, SQL migration, and all WRITE/POST endpoints.\n'
         'Block, unblock, renew, SMS send, temporary access.\n'
         'All endpoints require valid token (except login).',
         'PHP / MySQL\nDaloRADIUS codebase'],
    ]
)

doc.add_paragraph()

# ============ INTEGRATION FLOW ============
doc.add_heading('Integration Flow', level=1)
p = doc.add_paragraph()
p.add_run('Authentication Flow:\n').bold = True
p.add_run('1. Wael (ERP) calls POST /erp_login.php with username + password\n')
p.add_run('2. API validates against operators table, returns token (expires in 8 hours)\n')
p.add_run('3. Wael stores token and sends it on ALL subsequent requests:\n')
p.add_run('   Header: Authorization: Bearer <token>\n')
p.add_run('4. If token expires, API returns 401 -> ERP redirects to login\n\n')

p.add_run('API Flow:\n').bold = True
p.add_run('ERP Login Page (Wael)    -->  POST /erp_login (Laith)  -->  Token\n')
p.add_run('ERP Dashboard (Wael)    -->  GET + Bearer Token  -->  DaloRADIUS API (Hisham)  -->  MySQL/RADIUS\n')
p.add_run('ERP Action Buttons (Wael)  -->  POST + Bearer Token  -->  DaloRADIUS API (Laith)  -->  MySQL/RADIUS\n\n')

p.add_run('API Base URL: ').bold = True
p.add_run('http://<daloradius-server>/api/\n')
p.add_run('Auth Header: ').bold = True
p.add_run('Authorization: Bearer <token>\n')
p.add_run('Only /erp_login.php is public (no token needed)').italic = True

doc.add_paragraph()

# ============ LAITH TASKS ============
doc.add_heading('Laith Tasks - DB + Login/Auth + POST Endpoints', level=1)

add_styled_table(doc,
    ['#', 'Task', 'Assignee', 'Priority', 'Status', 'Notes'],
    [
        ['1', 'SQL migration (11_erp_api_tables.sql)\n'
              'Create tables:\n'
              '- api_tokens (id, operator_id, username, token,\n'
              '  expires_at, created_at, ip_address, is_revoked)\n'
              '- sms_notifications (id, username, phone, message,\n'
              '  sms_type, status, sent_at, created_at)\n'
              '- temporary_access (id, username, granted_by,\n'
              '  days_granted, start_date, expiry_date, status)',
         'Laith', 'HIGH', '[ ]', 'Must be done first\nbefore any other task'],

        ['2', 'erp_login.php\n'
              'POST {username, password}\n'
              'Validates against operators table\n'
              'Generates token, saves to api_tokens table\n'
              'Token expires in 8 hours\n'
              'Response: {success, token, expires_at,\n'
              '  operator: {id, username, firstname, lastname}}\n'
              'Logs login to system_action_log\n'
              'NO TOKEN REQUIRED (public endpoint)',
         'Laith', 'HIGH', '[ ]', 'Must be done first\nHisham & Wael depend on this'],

        ['3', 'erp_logout.php\n'
              'POST - Revokes current token\n'
              'Sets is_revoked=1 in api_tokens table\n'
              'Response: {success, message}\n'
              'REQUIRES TOKEN',
         'Laith', 'HIGH', '[ ]', ''],

        ['4', 'erp_refresh_token.php\n'
              'POST - Refreshes token before it expires\n'
              'Returns new token, revokes old one\n'
              'Only works if current token has < 2 hours remaining\n'
              'Response: {success, token, expires_at}\n'
              'REQUIRES TOKEN',
         'Laith', 'MEDIUM', '[ ]', ''],

        ['5', 'Update auth.php - Token validation\n'
              'Add token validation logic:\n'
              '- Check Authorization: Bearer header\n'
              '- Validate token exists in api_tokens table\n'
              '- Check not expired and not revoked\n'
              '- Set $authenticated_operator (id, username)\n'
              '  for use in all endpoints\n'
              '- Return 401 if invalid/expired with\n'
              '  error_code: TOKEN_EXPIRED or TOKEN_INVALID',
         'Laith', 'HIGH', '[ ]', 'All endpoints depend on this'],

        ['6', 'erp_block_user.php\n'
              'POST {username, reason}\n'
              'Adds user to block_user RADIUS group\n'
              'Uses RadiusAccessManager::revokeAccess()\n'
              'Logs action to system_action_log\n'
              'with operator name from token\n'
              'REQUIRES TOKEN',
         'Laith', 'HIGH', '[ ]', ''],

        ['7', 'erp_unblock_user.php\n'
              'POST {username}\n'
              'Removes from block_user group\n'
              'Restores plan access using\n'
              'RadiusAccessManager::reactivateUser()\n'
              'Logs action to system_action_log\n'
              'REQUIRES TOKEN',
         'Laith', 'HIGH', '[ ]', ''],

        ['8', 'erp_renew_bundle.php\n'
              'POST {username}\n'
              'Renews current bundle for full period\n'
              'Uses BundleManager::renewBundle()\n'
              'Deducts cost from balance, extends expiry\n'
              'created_by = operator username from token\n'
              'REQUIRES TOKEN',
         'Laith', 'HIGH', '[ ]', ''],

        ['9', 'erp_renew_temporary.php\n'
              'POST {username, days (default 2)}\n'
              'Grants temporary access for N days\n'
              'Saves to temporary_access table\n'
              'Sets RADIUS Expiration attribute\n'
              'Removes from block_user if blocked\n'
              'granted_by = operator username from token\n'
              'REQUIRES TOKEN',
         'Laith', 'HIGH', '[ ]', ''],

        ['10', 'erp_send_sms.php\n'
               'POST {username, message, sms_type}\n'
               'Queues SMS to sms_notifications table\n'
               'Types: expiry_warning, payment_reminder,\n'
               'block_notice, custom\n'
               'REQUIRES TOKEN',
         'Laith', 'MEDIUM', '[ ]', ''],
    ]
)

doc.add_paragraph()

# ============ HISHAM TASKS ============
doc.add_heading('Hisham Tasks - DaloRADIUS API (GET Endpoints)', level=1)

p = doc.add_paragraph()
p.add_run('All endpoints below REQUIRE a valid token in the Authorization header.\n').italic = True
p.add_run('Hisham must wait for Laith to complete tasks #1, #2, #5 (SQL + login + auth) before testing.').italic = True

doc.add_paragraph()

add_styled_table(doc,
    ['#', 'Task', 'Assignee', 'Priority', 'Status', 'Notes'],
    [
        ['11', 'erp_dashboard_summary.php\n'
               'GET - Dashboard aggregate counts\n'
               'Response: online_users, disconnected_users,\n'
               'blocked_users, upcoming_blocks_3days,\n'
               'upcoming_blocks_7days, total_users,\n'
               'active_bundles, expired_bundles,\n'
               'total_traffic_today_gb\n'
               'REQUIRES TOKEN',
         'Hisham', 'HIGH', '[ ]', ''],

        ['12', 'erp_online_users.php\n'
               'GET - Returns all currently online users\n'
               'Query: radacct WHERE acctstoptime IS NULL\n'
               'Response: username, fullname, NAS,\n'
               'session duration, download/upload MB,\n'
               'framed IP, plan, phone\n'
               'REQUIRES TOKEN',
         'Hisham', 'HIGH', '[ ]', ''],

        ['13', 'erp_blocked_users.php\n'
               'GET - Returns all blocked/disabled users\n'
               'Query: radusergroup WHERE groupname\n'
               'IN (block_user, Disabled-Users)\n'
               'Response: username, fullname, phone,\n'
               'block reason, plan, last session, balance\n'
               'REQUIRES TOKEN',
         'Hisham', 'HIGH', '[ ]', ''],

        ['14', 'erp_upcoming_blocks.php\n'
               'GET ?days=3\n'
               'Users whose bundle expires within N days\n'
               'Query: userbillinfo WHERE bundle_expiry_date\n'
               'BETWEEN NOW and NOW+N\n'
               'Response: username, fullname, phone, plan,\n'
               'expiry date, days remaining, balance,\n'
               'can_auto_renew\n'
               'REQUIRES TOKEN',
         'Hisham', 'HIGH', '[ ]', ''],

        ['15', 'erp_disconnected_users.php\n'
               'GET ?hours=24\n'
               'Recently disconnected users\n'
               'Query: radacct WHERE acctstoptime > NOW - N hours\n'
               'AND username NOT IN (online users)\n'
               'Response: username, fullname,\n'
               'last session end, last NAS, plan\n'
               'REQUIRES TOKEN',
         'Hisham', 'MEDIUM', '[ ]', ''],

        ['16', 'erp_user_detail.php\n'
               'GET ?username=X - Complete user profile\n'
               'Joins: userinfo + userbillinfo + radacct\n'
               '+ radusergroup + radgroupreply\n'
               'Response sections:\n'
               'personal{}, billing{}, speed_profile{},\n'
               'status{}, usage_summary{}\n'
               'REQUIRES TOKEN',
         'Hisham', 'HIGH', '[ ]', ''],

        ['17', 'erp_user_usage.php\n'
               'GET ?username=X&period=daily&month=2026-03\n'
               'OR ?username=X&period=monthly\n'
               'Query: radacct GROUP BY DATE or MONTH\n'
               'Response: per-day or per-month breakdown\n'
               '(download, upload, sessions, time)\n'
               '+ summary totals\n'
               'Supports pagination: ?page=1&limit=31\n'
               'REQUIRES TOKEN',
         'Hisham', 'HIGH', '[ ]', ''],

        ['18', 'erp_user_plan_profile.php\n'
               'GET ?username=X - Plan + speed profile\n'
               'Joins: userbillinfo + billing_plans\n'
               '+ billing_plans_profiles + radgroupreply\n'
               'Response: plan name, cost, type,\n'
               'bundle validity, RADIUS groups,\n'
               'Mikrotik-Rate-Limit, Session-Timeout\n'
               'REQUIRES TOKEN',
         'Hisham', 'MEDIUM', '[ ]', ''],

        ['19', 'erp_sms_history.php\n'
               'GET ?username=X\n'
               'SMS notifications sent to user\n'
               'Query: sms_notifications WHERE username=X\n'
               'Response: id, phone, message,\n'
               'sms_type, status, sent_at\n'
               'REQUIRES TOKEN',
         'Hisham', 'MEDIUM', '[ ]', ''],
    ]
)

doc.add_paragraph()

# ============ WAEL TASKS ============
doc.add_heading('Wael Tasks - ERP Frontend (Consumes All APIs)', level=1)

p = doc.add_paragraph()
p.add_run('Wael must implement login flow first.\n').bold = True
p.add_run('After login, store the token and send it with every API call:\n')
p.add_run('Authorization: Bearer <token>\n').bold = True
p.add_run('If any API returns 401 (TOKEN_EXPIRED), redirect user to login page.\n')
p.add_run('If token is close to expiry, call POST /erp_refresh_token.php to get new token.').italic = True

doc.add_paragraph()

add_styled_table(doc,
    ['#', 'Task', 'Assignee', 'APIs Used', 'Priority', 'Status', 'Notes'],
    [
        ['20', 'Login Page\n'
               'Username + password form\n'
               'Call POST /erp_login\n'
               'Store token (localStorage/cookie)\n'
               'Redirect to dashboard on success\n'
               'Show error message on failure',
         'Wael', 'POST erp_login\n(built by Laith)', 'HIGH', '[ ]', 'Must be done first'],

        ['21', 'Auth Middleware / HTTP Client\n'
               'Create shared HTTP client that:\n'
               '- Adds Authorization: Bearer header\n'
               '  to all requests\n'
               '- Intercepts 401 responses ->\n'
               '  redirect to login\n'
               '- Auto-refreshes token when\n'
               '  close to expiry\n'
               '- Stores/retrieves token from storage',
         'Wael', 'POST erp_refresh_token\nPOST erp_logout\n(built by Laith)', 'HIGH', '[ ]', 'Must be done first\nAll pages use this'],

        ['22', 'Logout Button\n'
               'Button in top navbar\n'
               'Calls POST /erp_logout\n'
               'Clears stored token\n'
               'Redirects to login page',
         'Wael', 'POST erp_logout\n(built by Laith)', 'HIGH', '[ ]', ''],

        ['23', 'Dashboard Page\n'
               '4 summary cards:\n'
               'Online Users, Disconnected,\n'
               'Blocked, Upcoming Blocks\n'
               'Each card shows count from API\n'
               'Auto-refresh every 60 seconds',
         'Wael', 'GET erp_dashboard_summary\n(built by Hisham)', 'HIGH', '[ ]', ''],

        ['24', 'Online Users Table\n'
               'Columns: username, name, NAS,\n'
               'duration, download, upload, IP\n'
               'Click row -> user detail page\n'
               'Auto-refresh every 30 seconds',
         'Wael', 'GET erp_online_users\n(built by Hisham)', 'HIGH', '[ ]', ''],

        ['25', 'Blocked Users Table\n'
               'Columns: username, name, phone,\n'
               'reason, plan, last session\n'
               'Unblock button on each row\n'
               'Click row -> user detail page',
         'Wael', 'GET erp_blocked_users (Hisham)\nPOST erp_unblock_user (Laith)', 'HIGH', '[ ]', ''],

        ['26', 'Upcoming Blocks Table\n'
               'Columns: username, name, phone,\n'
               'plan, expiry, days left, balance\n'
               'Renew + Renew 2 Days buttons\n'
               'Click row -> user detail page',
         'Wael', 'GET erp_upcoming_blocks (Hisham)\nPOST erp_renew_bundle (Laith)\nPOST erp_renew_temporary (Laith)', 'HIGH', '[ ]', ''],

        ['27', 'Disconnected Users Table\n'
               'Columns: username, name,\n'
               'last session, NAS, plan\n'
               'Click row -> user detail page',
         'Wael', 'GET erp_disconnected_users\n(built by Hisham)', 'MEDIUM', '[ ]', ''],

        ['28', 'User Detail Page\n'
               'Full profile page for single user:\n'
               '- Personal info card\n'
               '- Billing info (plan, balance, bundle)\n'
               '- Speed profile (download/upload)\n'
               '- Online status indicator\n'
               '- Usage summary',
         'Wael', 'GET erp_user_detail (Hisham)\nGET erp_user_plan_profile (Hisham)', 'HIGH', '[ ]', ''],

        ['29', 'Usage Charts Page\n'
               'Daily traffic chart (bar/line)\n'
               'Monthly traffic chart\n'
               'Date range / month selector\n'
               'Download/Upload breakdown',
         'Wael', 'GET erp_user_usage (Hisham)\n?period=daily\n?period=monthly', 'HIGH', '[ ]', ''],

        ['30', 'Block User Button\n'
               'Button on user detail + tables\n'
               'Confirmation dialog with reason\n'
               'Success/error feedback\n'
               'Refresh table after action',
         'Wael', 'POST erp_block_user\n(built by Laith)', 'HIGH', '[ ]', ''],

        ['31', 'Unblock User Button\n'
               'Button on blocked table + detail\n'
               'Confirmation dialog\n'
               'Refreshes table after success',
         'Wael', 'POST erp_unblock_user\n(built by Laith)', 'HIGH', '[ ]', ''],

        ['32', 'Renew Bundle Button\n'
               'Shows plan cost and balance\n'
               'Confirmation before renewal\n'
               'Shows new expiry after success',
         'Wael', 'POST erp_renew_bundle\n(built by Laith)', 'HIGH', '[ ]', ''],

        ['33', 'Renew 2 Days Button\n'
               'Input field for days (default 2)\n'
               'Confirmation dialog\n'
               'Shows temporary expiry after success',
         'Wael', 'POST erp_renew_temporary\n(built by Laith)', 'HIGH', '[ ]', ''],

        ['34', 'SMS Log Page\n'
               'Table of sent SMS per user\n'
               'Columns: date, phone, message,\n'
               'type, status\n'
               'Filter by type/status',
         'Wael', 'GET erp_sms_history\n(built by Hisham)', 'MEDIUM', '[ ]', ''],

        ['35', 'Send SMS Button\n'
               'Message input + type selector\n'
               'Delivery status feedback\n'
               'Types: expiry warning,\n'
               'payment reminder,\n'
               'block notice, custom',
         'Wael', 'POST erp_send_sms\n(built by Laith)', 'MEDIUM', '[ ]', ''],

        ['36', 'Final Integration & Testing\n'
               'Test all API calls end-to-end\n'
               'Error handling for failures\n'
               'Loading states for all tables\n'
               'Token expiry handling\n'
               'Responsive design check',
         'Wael', 'ALL endpoints', 'HIGH', '[ ]', ''],
    ]
)

doc.add_paragraph()

# ============ TASK SUMMARY ============
doc.add_heading('Task Summary', level=1)

add_styled_table(doc,
    ['Developer', 'Role', 'Tasks', 'Total'],
    [
        ['Laith', 'DB + Login/Auth + POST endpoints', '#1 - #10', '10 tasks'],
        ['Hisham', 'GET endpoints', '#11 - #19', '9 tasks'],
        ['Wael', 'ERP Frontend (consumes all APIs)', '#20 - #36', '17 tasks'],
    ]
)

doc.add_paragraph()

# ============ DEPENDENCY ORDER ============
doc.add_heading('Task Dependencies (Build Order)', level=1)

p = doc.add_paragraph()
p.add_run('These tasks MUST be completed in order:\n\n').bold = True
p.add_run('1. Laith: Task #1 (SQL migration) - creates all new tables\n')
p.add_run('2. Laith: Task #2 (erp_login.php) - login endpoint\n')
p.add_run('3. Laith: Task #5 (auth.php update) - token validation\n')
p.add_run('   --> After these 3, Hisham and Wael can start their work\n\n')
p.add_run('4. Wael: Task #20 (Login page) + Task #21 (Auth middleware)\n')
p.add_run('   --> After these, Wael can build any dashboard page\n\n')
p.add_run('5. All other tasks can be done in any order (in parallel)')

doc.add_paragraph()


# ============ SIGN-OFF ============
doc.add_heading('Sign-Off', level=1)

add_styled_table(doc,
    ['Developer', 'Signature', 'Date'],
    [
        ['Wael (ERP Frontend)', '', '____/____/2026'],
        ['Hisham (API GET Endpoints)', '', '____/____/2026'],
        ['Laith (API DB + POST + Auth)', '', '____/____/2026'],
    ]
)

# Save
output_path = os.path.join(os.path.dirname(__file__), 'ERP_Integration_Tasks.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
